"""
Behavior-organized tests for telechat_pkg.document_extract.

This module turns uploaded files (PDF / DOCX / CSV / TXT / code) into text.
Optional deps:
    - PyMuPDF (imported as ``fitz``)  -> NOT installed in this env.
    - python-docx (imported as ``docx``) -> installed in this env.

So the PDF "real extraction" tests inject a *fake* ``fitz`` module into
``sys.modules`` (the genuine not-installed path is exercised directly), while
the DOCX not-installed path is simulated by hiding the real ``docx`` module.

Run:
    COVERAGE_FILE=/tmp/.cov_0014 python -m pytest -q tests/test_document_extract.py \
        --cov=telechat_pkg.document_extract --cov-report=term-missing
"""

from __future__ import annotations

import builtins
import importlib
import os
import sys

import pytest

os.environ.setdefault("TELEGRAM_BOT_TOKEN", "test-token")
os.environ.setdefault("ANTHROPIC_API_KEY", "test-api-key")

from telechat_pkg import document_extract as de
from telechat_pkg.document_extract import (
    ExtractResult,
    available_formats,
    extract,
    extract_csv,
    extract_docx,
    extract_pdf,
    extract_text_file,
    summarize_extraction,
)


# ──────────────────────────────────────────────────────────────────────────────
# Fakes / helpers
# ──────────────────────────────────────────────────────────────────────────────


class _FakePage:
    def __init__(self, text: str):
        self._text = text

    def get_text(self) -> str:
        return self._text


class _FakeDoc:
    """Minimal stand-in for a fitz.Document."""

    def __init__(self, page_texts):
        self._pages = [_FakePage(t) for t in page_texts]
        self.closed = False

    def __len__(self):
        return len(self._pages)

    def __getitem__(self, i):
        return self._pages[i]

    def close(self):
        self.closed = True


class _FakeFitz:
    """A fake ``fitz`` module. ``open`` returns whatever the test queued."""

    def __init__(self):
        self._next = None
        self._raise = None
        self.last_path = None
        self.opened_doc = None

    def queue_pages(self, page_texts):
        self._next = page_texts
        self._raise = None

    def queue_error(self, exc):
        self._raise = exc
        self._next = None

    def open(self, path):
        self.last_path = path
        if self._raise is not None:
            raise self._raise
        doc = _FakeDoc(self._next or [])
        self.opened_doc = doc
        return doc


@pytest.fixture
def fake_fitz(monkeypatch):
    """Inject a fake ``fitz`` so the installed-path of extract_pdf runs."""
    fake = _FakeFitz()
    monkeypatch.setitem(sys.modules, "fitz", fake)
    return fake


@pytest.fixture
def hide_module(monkeypatch):
    """Return a function that makes ``import <name>`` raise ImportError."""

    def _hide(name: str):
        monkeypatch.setitem(sys.modules, name, None)
        real_import = builtins.__import__

        def fake_import(modname, *args, **kwargs):
            if modname == name or modname.startswith(name + "."):
                raise ImportError(f"hidden: {modname}")
            return real_import(modname, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", fake_import)

    return _hide


def _write(tmp_path, name, data, *, binary=False):
    p = tmp_path / name
    if binary:
        p.write_bytes(data)
    else:
        p.write_text(data, encoding="utf-8")
    return str(p)


def _make_docx(tmp_path, name, paragraphs=(), table_rows=None):
    """Build a real .docx on disk using python-docx (installed here)."""
    import docx

    doc = docx.Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_rows:
        ncols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=ncols)
        for ri, row in enumerate(table_rows):
            for ci, val in enumerate(row):
                table.rows[ri].cells[ci].text = val
    path = str(tmp_path / name)
    doc.save(path)
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 1. PDF extraction (PyMuPDF/fitz path)
# ──────────────────────────────────────────────────────────────────────────────


class TestPdfExtraction:
    def test_pymupdf_not_installed_returns_error_result(self):
        # fitz is genuinely absent in this env: the real not-installed path.
        result = extract_pdf("anything.pdf")
        assert result.format == "pdf"
        assert result.text == ""
        assert result.pages == 0
        assert result.error is not None
        assert "PyMuPDF" in result.error

    def test_extract_two_pages(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["Hello from page one.", "Second page text."])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert result.error is None
        assert result.format == "pdf"
        assert result.pages == 2
        assert "Hello from page one." in result.text
        assert "Second page text." in result.text

    def test_doc_is_closed(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["content"])
        path = _write(tmp_path, "doc.pdf", "stub")
        extract_pdf(path)
        assert fake_fitz.opened_doc.closed is True

    def test_blank_pages_skipped(self, fake_fitz, tmp_path):
        # Whitespace-only pages contribute neither text nor page count.
        fake_fitz.queue_pages(["   \n  ", "real content", "\t"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert result.pages == 1
        assert "real content" in result.text

    def test_open_raises_returns_error(self, fake_fitz, tmp_path):
        fake_fitz.queue_error(RuntimeError("corrupt pdf"))
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert result.text == ""
        assert result.pages == 0
        assert result.error == "corrupt pdf"

    def test_error_message_truncated_to_200(self, fake_fitz, tmp_path):
        fake_fitz.queue_error(RuntimeError("x" * 500))
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert len(result.error) == 200

    def test_truncation_when_text_too_long(self, fake_fitz, tmp_path, monkeypatch):
        monkeypatch.setattr(de, "MAX_TEXT_LENGTH", 50)
        fake_fitz.queue_pages(["A" * 200])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert result.truncated is True
        assert result.text.endswith("[...truncated...]")

    def test_no_truncation_for_short_text(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["short"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert result.truncated is False


# ──────────────────────────────────────────────────────────────────────────────
# 2. DOCX extraction (python-docx path)
# ──────────────────────────────────────────────────────────────────────────────


class TestDocxExtraction:
    def test_python_docx_not_installed_returns_error_result(self, hide_module):
        hide_module("docx")
        result = extract_docx("anything.docx")
        assert result.format == "docx"
        assert result.text == ""
        assert result.pages == 0
        assert result.error is not None
        assert "python-docx" in result.error

    def test_extract_paragraphs(self, tmp_path):
        path = _make_docx(tmp_path, "d.docx", paragraphs=["First para", "Second para"])
        result = extract_docx(path)
        assert result.error is None
        assert result.format == "docx"
        assert "First para" in result.text
        assert "Second para" in result.text
        assert result.pages == 2  # two sections

    def test_empty_paragraphs_skipped(self, tmp_path):
        path = _make_docx(tmp_path, "d.docx", paragraphs=["kept", "   ", ""])
        result = extract_docx(path)
        assert result.pages == 1
        assert "kept" in result.text

    def test_table_cells_extracted(self, tmp_path):
        path = _make_docx(
            tmp_path,
            "d.docx",
            paragraphs=["intro"],
            table_rows=[["a", "b"], ["c", "d"]],
        )
        result = extract_docx(path)
        assert "a | b" in result.text
        assert "c | d" in result.text

    def test_corrupt_docx_returns_error(self, tmp_path):
        # A non-zip file is not a valid .docx -> docx.Document raises.
        path = _write(tmp_path, "bad.docx", "not a real docx file")
        result = extract_docx(path)
        assert result.text == ""
        assert result.error is not None

    def test_truncation_when_text_too_long(self, tmp_path, monkeypatch):
        monkeypatch.setattr(de, "MAX_TEXT_LENGTH", 20)
        path = _make_docx(tmp_path, "d.docx", paragraphs=["B" * 200])
        result = extract_docx(path)
        assert result.truncated is True
        assert result.text.endswith("[...truncated...]")


# ──────────────────────────────────────────────────────────────────────────────
# 3. CSV extraction
# ──────────────────────────────────────────────────────────────────────────────


class TestCsvExtraction:
    def test_basic_csv(self, tmp_path):
        path = _write(tmp_path, "data.csv", "name,age\nAlice,30\nBob,25\n")
        result = extract_csv(path)
        assert result.error is None
        assert result.format == "csv"
        assert "name | age" in result.text
        assert "Alice | 30" in result.text

    def test_csv_row_count(self, tmp_path):
        path = _write(tmp_path, "data.csv", "a,b\n1,2\n3,4\n")
        result = extract_csv(path)
        assert result.pages == 3  # header + 2 data rows

    def test_csv_sniff_fallback_to_excel(self, tmp_path):
        # Single-column content gives the Sniffer nothing to detect -> excel.
        path = _write(tmp_path, "data.csv", "onlyonecolumn\nrow2\nrow3\n")
        result = extract_csv(path)
        assert result.error is None
        assert "onlyonecolumn" in result.text

    def test_csv_open_error(self, tmp_path):
        result = extract_csv(str(tmp_path / "does_not_exist.csv"))
        assert result.text == ""
        assert result.error is not None

    def test_csv_row_limit(self, tmp_path, monkeypatch):
        # Many rows but keep MAX_TEXT_LENGTH big so we hit the 10000-row cap,
        # not the char cap.
        monkeypatch.setattr(de, "MAX_TEXT_LENGTH", 10_000_000)
        rows = "\n".join(f"r{i},v{i}" for i in range(10050))
        path = _write(tmp_path, "big.csv", rows + "\n")
        result = extract_csv(path)
        assert "[...truncated at 10000 rows...]" in result.text

    def test_csv_char_truncation(self, tmp_path, monkeypatch):
        monkeypatch.setattr(de, "MAX_TEXT_LENGTH", 30)
        path = _write(tmp_path, "data.csv", "col\n" + "\n".join("x" * 10 for _ in range(20)))
        result = extract_csv(path)
        assert result.truncated is True
        assert result.text.endswith("[...truncated...]")


# ──────────────────────────────────────────────────────────────────────────────
# 4. Plain text / code files
# ──────────────────────────────────────────────────────────────────────────────


class TestTextFileExtraction:
    def test_plain_text(self, tmp_path):
        path = _write(tmp_path, "notes.txt", "line one\nline two\nline three")
        result = extract_text_file(path)
        assert result.error is None
        assert result.format == "txt"
        assert result.text == "line one\nline two\nline three"

    def test_line_count_reported_as_pages(self, tmp_path):
        path = _write(tmp_path, "notes.txt", "a\nb\nc")
        result = extract_text_file(path)
        assert result.pages == 3  # count("\n") + 1

    def test_code_file_format_from_extension(self, tmp_path):
        path = _write(tmp_path, "script.py", "print('hi')\n")
        result = extract_text_file(path)
        assert result.format == "py"

    def test_no_extension_defaults_to_txt(self, tmp_path):
        path = _write(tmp_path, "README", "some text")
        result = extract_text_file(path)
        assert result.format == "txt"

    def test_truncation(self, tmp_path, monkeypatch):
        monkeypatch.setattr(de, "MAX_TEXT_LENGTH", 10)
        path = _write(tmp_path, "big.txt", "Z" * 100)
        result = extract_text_file(path)
        assert result.truncated is True
        assert result.text.endswith("[...truncated...]")

    def test_read_error_returns_error_result(self, tmp_path, monkeypatch):
        def boom(*a, **k):
            raise OSError("disk gone")

        monkeypatch.setattr(builtins, "open", boom)
        result = extract_text_file(str(tmp_path / "x.txt"))
        assert result.text == ""
        assert result.format == "txt"
        assert result.error == "disk gone"


# ──────────────────────────────────────────────────────────────────────────────
# 5. Encoding fallback (errors="replace")
# ──────────────────────────────────────────────────────────────────────────────


class TestEncodingFallback:
    def test_invalid_utf8_bytes_are_replaced_not_crashed(self, tmp_path):
        # 0xFF / 0xFE are invalid as UTF-8; module opens with errors="replace".
        path = _write(tmp_path, "weird.txt", b"valid text \xff\xfe more text", binary=True)
        result = extract_text_file(path)
        assert result.error is None
        assert "valid text" in result.text
        assert "more text" in result.text
        # The replacement character appears instead of crashing.
        assert "�" in result.text

    def test_latin1_bytes_via_csv(self, tmp_path):
        # Latin-1 encoded accented chars -> replaced, no UnicodeDecodeError.
        path = _write(tmp_path, "data.csv", b"label,city\ncost,caf\xe9\n", binary=True)
        result = extract_csv(path)
        assert result.error is None
        assert "label | city" in result.text
        # The undecodable byte became a replacement char rather than crashing.
        assert "�" in result.text

    def test_extract_routes_bad_encoding_file(self, tmp_path):
        path = _write(tmp_path, "log.log", b"\xff\xfe\xfacontent", binary=True)
        result = extract(path)
        assert result.error is None
        assert "content" in result.text


# ──────────────────────────────────────────────────────────────────────────────
# 6. Format detection / dispatch via extract()
# ──────────────────────────────────────────────────────────────────────────────


class TestExtractDispatch:
    def test_file_not_found(self, tmp_path):
        result = extract(str(tmp_path / "nope.txt"))
        assert result.format == "unknown"
        assert "File not found" in result.error

    def test_empty_file(self, tmp_path):
        path = _write(tmp_path, "empty.txt", "")
        result = extract(path)
        assert result.format == "empty"
        assert result.error == "File is empty"

    def test_file_too_large(self, tmp_path, monkeypatch):
        monkeypatch.setattr(de, "MAX_FILE_SIZE", 5)
        path = _write(tmp_path, "big.txt", "way more than five bytes")
        result = extract(path)
        assert result.format == "unknown"
        assert "File too large" in result.error

    def test_dispatch_pdf(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["pdf body"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract(path)
        assert result.format == "pdf"
        assert "pdf body" in result.text

    def test_dispatch_docx(self, tmp_path):
        path = _make_docx(tmp_path, "d.docx", paragraphs=["docx body"])
        result = extract(path)
        assert result.format == "docx"
        assert "docx body" in result.text

    def test_dispatch_csv(self, tmp_path):
        path = _write(tmp_path, "data.csv", "a,b\n1,2\n")
        result = extract(path)
        assert result.format == "csv"

    def test_dispatch_code_file(self, tmp_path):
        path = _write(tmp_path, "main.go", "package main\n")
        result = extract(path)
        assert result.format == "go"

    def test_dispatch_txt(self, tmp_path):
        path = _write(tmp_path, "notes.txt", "hello")
        result = extract(path)
        assert result.format == "txt"

    def test_dispatch_md(self, tmp_path):
        path = _write(tmp_path, "readme.md", "# title")
        result = extract(path)
        assert result.format == "md"


# ──────────────────────────────────────────────────────────────────────────────
# 7. Unsupported / unknown extensions
# ──────────────────────────────────────────────────────────────────────────────


class TestUnsupportedFormat:
    def test_unknown_extension_falls_back_to_text(self, tmp_path):
        # Unknown extensions are read as text anyway (the `else` branch).
        path = _write(tmp_path, "thing.xyz", "readable as text")
        result = extract(path)
        assert result.error is None
        assert "readable as text" in result.text
        # format derived from extract_text_file => the extension itself.
        assert result.format == "xyz"

    def test_no_extension_falls_back_to_text(self, tmp_path):
        path = _write(tmp_path, "noext", "still text")
        result = extract(path)
        assert result.error is None
        assert "still text" in result.text

    def test_binary_unknown_extension_does_not_crash(self, tmp_path):
        # errors="replace" means even a "binary" unknown file decodes.
        path = _write(tmp_path, "blob.bin", b"\x00\x01\x02hello", binary=True)
        result = extract(path)
        # The fallback path reads it as text rather than erroring.
        assert result.error is None
        assert "hello" in result.text

    # NOTE: extract()'s unknown-extension branch wraps extract_text_file() in a
    # try/except returning an "Unsupported file format" error (source lines
    # 247-248). That except is unreachable dead code: extract_text_file()
    # already catches every exception internally and always returns an
    # ExtractResult, so the outer try never sees a raise. Left uncovered and
    # documented rather than contorting a test to fake it.


# ──────────────────────────────────────────────────────────────────────────────
# 8. Page / section boundary preservation
# ──────────────────────────────────────────────────────────────────────────────


class TestPageBoundaries:
    def test_pdf_page_markers_present(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["alpha", "beta", "gamma"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert "--- Page 1 ---" in result.text
        assert "--- Page 2 ---" in result.text
        assert "--- Page 3 ---" in result.text

    def test_pdf_page_numbering_skips_blank(self, fake_fitz, tmp_path):
        # Page numbering reflects the *original* index, not the kept index.
        fake_fitz.queue_pages(["first", "   ", "third"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert "--- Page 1 ---" in result.text
        assert "--- Page 3 ---" in result.text
        assert "--- Page 2 ---" not in result.text
        assert result.pages == 2

    def test_pdf_pages_joined_with_blank_line(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["one", "two"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert "\n\n--- Page 2 ---" in result.text

    def test_pdf_page_content_follows_marker(self, fake_fitz, tmp_path):
        fake_fitz.queue_pages(["the body text"])
        path = _write(tmp_path, "doc.pdf", "stub")
        result = extract_pdf(path)
        assert "--- Page 1 ---\nthe body text" in result.text

    def test_docx_sections_joined_with_blank_line(self, tmp_path):
        path = _make_docx(tmp_path, "d.docx", paragraphs=["sec one", "sec two"])
        result = extract_docx(path)
        assert "sec one\n\nsec two" in result.text

    def test_csv_rows_joined_with_newline(self, tmp_path):
        path = _write(tmp_path, "data.csv", "a,b\n1,2\n")
        result = extract_csv(path)
        assert "a | b\n1 | 2" in result.text


# ──────────────────────────────────────────────────────────────────────────────
# 9. available_formats / _check_deps
# ──────────────────────────────────────────────────────────────────────────────


class TestAvailableFormats:
    def test_base_formats_always_present(self):
        fmts = available_formats()
        assert "txt" in fmts
        assert "csv" in fmts
        assert "code files" in fmts

    def test_docx_listed_when_installed(self):
        # python-docx is installed in this env.
        assert "docx" in available_formats()

    def test_pdf_absent_when_fitz_missing(self):
        # fitz is genuinely not installed here.
        assert "pdf" not in available_formats()

    def test_pdf_listed_when_fitz_present(self, fake_fitz):
        assert "pdf" in available_formats()

    def test_check_deps_reports_docx_true_fitz_false(self):
        deps = de._check_deps()
        assert deps["docx"] is True
        assert deps["fitz"] is False

    def test_check_deps_fitz_true_when_injected(self, fake_fitz):
        deps = de._check_deps()
        assert deps["fitz"] is True

    def test_check_deps_false_when_hidden(self, hide_module):
        hide_module("docx")
        deps = de._check_deps()
        assert deps["docx"] is False


# ──────────────────────────────────────────────────────────────────────────────
# 10. summarize_extraction display formatting
# ──────────────────────────────────────────────────────────────────────────────


class TestSummarize:
    def test_error_summary(self):
        r = ExtractResult(text="", pages=0, format="pdf", error="boom")
        out = summarize_extraction(r)
        assert "Error extracting pdf" in out
        assert "boom" in out

    def test_success_summary_header(self):
        r = ExtractResult(text="hello world", pages=3, format="txt")
        out = summarize_extraction(r)
        assert "Extracted from TXT (3 pages/sections)" in out
        assert "hello world" in out

    def test_truncated_note(self):
        r = ExtractResult(text="x", pages=1, format="txt", truncated=True)
        out = summarize_extraction(r)
        assert "(truncated due to size)" in out

    def test_preview_capped_with_ellipsis(self):
        r = ExtractResult(text="A" * 600, pages=1, format="txt")
        out = summarize_extraction(r)
        assert out.endswith("...")
        # 500-char preview plus the ellipsis.
        assert "A" * 500 in out

    def test_short_text_no_ellipsis(self):
        r = ExtractResult(text="short", pages=1, format="txt")
        out = summarize_extraction(r)
        assert not out.rstrip().endswith("...")
